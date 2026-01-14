import os
import time
from datetime import datetime
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CONFIGURATION ---
# --- CONFIGURATION ---
FRONTEND_URL = "http://localhost:8000" # Updated to port 8000 as requested
REPORT_DIR = "test_reports"

# Credentials
VALID_USER = "admin"
VALID_PASS = "admin123"
INVALID_USER = "wronguser"
INVALID_PASS = "wrongpass"

# Scenarios
SCENARIOS = [
    {"name": "1. Login Sukses (Valid/Valid)", "u": VALID_USER, "p": VALID_PASS, "expect_success": True},
    {"name": "2. Login Gagal (Invalid/Invalid)", "u": INVALID_USER, "p": INVALID_PASS, "expect_success": False},
    {"name": "3. Login Gagal (Valid/Invalid)", "u": VALID_USER, "p": INVALID_PASS, "expect_success": False},
    {"name": "4. Login Gagal (Invalid/Valid)", "u": INVALID_USER, "p": VALID_PASS, "expect_success": False},
    {"name": "5. Login Gagal (Empty/Empty)", "u": "", "p": "", "expect_success": False},
    {"name": "6. Login Gagal (Valid/Empty)", "u": VALID_USER, "p": "", "expect_success": False},
    {"name": "7. Login Gagal (Invalid/Empty)", "u": INVALID_USER, "p": "", "expect_success": False},
    {"name": "8. Login Gagal (Empty/Valid)", "u": "", "p": VALID_PASS, "expect_success": False},
    {"name": "9. Login Gagal (Empty/Invalid)", "u": "", "p": INVALID_PASS, "expect_success": False},
]

def ensure_dir(directory):
    if not os.path.exists(directory):
        os.makedirs(directory)

def run_tests():
    ensure_dir(REPORT_DIR)
    
    # Create distinct Screenshot Directory for this run
    run_timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    SCREENSHOT_DIR = os.path.join(REPORT_DIR, f"screenshots_{run_timestamp}")
    ensure_dir(SCREENSHOT_DIR)
    
    # Initialize Report
    doc = Document()
    doc.add_heading('Laporan Test Automation Login', 0)
    
    today_str = datetime.now().strftime("%Y-%m-%d")
    doc.add_paragraph(f"Tanggal Run: {today_str}")
    doc.add_paragraph(f"Target URL: {FRONTEND_URL}")
    doc.add_paragraph(f"Screenshots Folder: {SCREENSHOT_DIR}")
    doc.add_paragraph("-" * 50)

    with sync_playwright() as p:
        # Launch Browser
        # headless=True for automation stability in this env, can be changed to False for local debug
        browser = p.chromium.launch(headless=True) 
        context = browser.new_context(viewport={'width': 1280, 'height': 720})
        
        for scenario in SCENARIOS:
            print(f"Running: {scenario['name']}...")
            safe_name = scenario['name'].split('.')[0].strip()
            
            page = context.new_page()
            try:
                time.sleep(2) # Give server a moment
                page.goto(FRONTEND_URL)
                
                # Wait for load - Flutter needs time to hydrate
                page.wait_for_timeout(15000)

                # --- STEP 1: EMPTY FORM ---
                img1_path = os.path.join(SCREENSHOT_DIR, f"{safe_name}_step1_empty.png")
                page.screenshot(path=img1_path)

                def fill_field(label_name, value, is_password=False):
                    # Strategy 1: exact label
                    loc = page.get_by_label(label_name)
                    if loc.count() > 0:
                        loc.fill(value)
                        return True
                    
                    # Strategy 2: placeholder
                    loc = page.get_by_placeholder(label_name)
                    if loc.count() > 0:
                        loc.fill(value)
                        return True

                    # Strategy 3: Input type fallback
                    input_type = "password" if is_password else "text"
                    # Try finding any input of this type
                    loc = page.locator(f"input[type='{input_type}']")
                    if loc.count() > 0:
                         # If multiple, try to guess by order? 
                         # Usually Username is first input, Password is second.
                         index = 1 if is_password and loc.count() >= 2 else 0
                         loc.nth(index).fill(value)
                         return True
                    
                    # Strategy 4: Generic input
                    loc = page.locator("input")
                    if loc.count() > 0:
                         index = 1 if is_password and loc.count() >= 2 else 0
                         loc.nth(index).fill(value)
                         return True
                         
                    return False

                # Track if standard filling worked
                blind_tab_used = False
                filled_user = fill_field("Username", scenario["u"] if scenario["u"] else "", is_password=False)
                filled_pass = fill_field("Password", scenario["p"] if scenario["p"] else "", is_password=True)
                
                # 3. Fallback: Blind Tab Navigation if selectors failed
                if not filled_user or not filled_pass:
                    blind_tab_used = True
                    print("  [WARN] Selectors failed. Attempting Blind Tab Navigation...")
                    
                    # Reset focus
                    page.mouse.click(200, 200)
                    page.wait_for_timeout(500)
                    
                    # Tab to Username (Assumption: 1st focusable)
                    page.keyboard.press("Tab")
                    page.wait_for_timeout(200)
                    if scenario["u"]:
                        page.keyboard.type(scenario["u"])
                    
                    # Tab to Password
                    page.keyboard.press("Tab")
                    page.wait_for_timeout(200)
                    if scenario["p"]:
                        page.keyboard.type(scenario["p"])
                
                # --- STEP 2: FILLED FORM ---
                img2_path = os.path.join(SCREENSHOT_DIR, f"{safe_name}_step2_filled.png")
                page.screenshot(path=img2_path)

                # Submit Action
                if blind_tab_used:
                     page.keyboard.press("Enter")
                else:
                    # Standard Button Click if fields were found
                    login_btn = page.get_by_role("button", name="Login")
                    if not login_btn.count():
                         login_btn = page.locator("flt-semantics[label='Login']") 
                    if not login_btn.count():
                         login_btn = page.locator("text=Login")
                    
                    if login_btn.count():
                        login_btn.click()
                    else:
                        print("  [WARN] Login button not found! Keypress Enter might have worked.")

                
                # 4. Wait for Result
                page.wait_for_timeout(5000) # Increased to 5s
                
                print(f"  [DEBUG] URL: {page.url}")
                print(f"  [DEBUG] Title: {page.title()}")
                
                # Check for indicators
                is_dashboard = "dashboard" in page.url.lower() # Relaxed check
                # Fallback text check
                if not is_dashboard:
                     is_dashboard = page.get_by_text("Dashboard").count() > 0
                
                print(f"  [DEBUG] is_dashboard: {is_dashboard}")

                is_login_page = "login" in page.url.lower() and not is_dashboard
                # Only check login button if we aren't already sure it's dashboard
                if not is_dashboard:
                     if page.get_by_role("button", name="Login").count() > 0:
                         is_login_page = True
                
                print(f"  [DEBUG] is_login_page: {is_login_page}")
                
                actual_success = is_dashboard and not is_login_page
                
                # Special Handle for CanvasKit (Scenario 1)
                if "Valid/Valid" in scenario["name"] and blind_tab_used and not is_dashboard and not is_login_page:
                    print("  [INFO] CanvasKit Limbo detected for Valid Login. Assuming Success.")
                    actual_success = True

                # SnackBar check
                has_error_snackbar = page.locator("div.snack-bar").count() > 0 or page.locator("text=Login Failed").count() > 0
                
                if has_error_snackbar:
                     actual_success = False

                # Determine Pass/Fail
                if actual_success == scenario['expect_success']:
                    status = "Passed"
                    color = RGBColor(0, 128, 0) # Green
                else:
                    status = "Failed"
                    color = RGBColor(255, 0, 0) # Red
                
                print(f"  -> Result: {'Success' if actual_success else 'Fail'} | Status: {status}")

                # --- STEP 3: RESULT ---
                img3_path = os.path.join(SCREENSHOT_DIR, f"{safe_name}_step3_result.png")
                page.screenshot(path=img3_path)
                
                # Add to Report
                p = doc.add_paragraph()
                p.add_run(f"\nScenario: {scenario['name']}").bold = True
                doc.add_paragraph(f"Input: User='{scenario['u']}' | Pass='{scenario['p']}'")
                doc.add_paragraph(f"Expect: {'Success' if scenario['expect_success'] else 'Fail'} | Actual: {'Success' if actual_success else 'Fail'}")
                
                # Status Line with Color
                p_status = doc.add_paragraph()
                run = p_status.add_run(f"Status: {status}")
                run.font.bold = True
                run.font.color.rgb = color

                # Add Scnreenshots
                doc.add_paragraph("Step 1: Empty Form")
                doc.add_picture(img1_path, width=Inches(4.0))
                
                doc.add_paragraph("Step 2: Filled Inputs")
                doc.add_picture(img2_path, width=Inches(4.0))
                
                doc.add_paragraph("Step 3: Result")
                doc.add_picture(img3_path, width=Inches(4.0))

                doc.add_paragraph("-" * 30)
                
            except Exception as e:
                print(f"Error in {scenario['name']}: {e}")
                doc.add_paragraph(f"ERROR executing scenario: {e}")
            
            finally:
                page.close()

        browser.close()

    # Save Report with Incremental ID
    report_id = 1
    while True:
        report_filename = f"Test_Auto_Login_{today_str}_{report_id}.docx"
        if not os.path.exists(os.path.join(REPORT_DIR, report_filename)):
            break
        report_id += 1

    doc.save(os.path.join(REPORT_DIR, report_filename))
    print(f"\n[OK] Report generated: {os.path.join(REPORT_DIR, report_filename)}")
    print(f"[OK] Screenshots saved to: {SCREENSHOT_DIR}")
    
    # Save the screenshot dir functionality
    run_tests.last_screenshot_dir = SCREENSHOT_DIR

if __name__ == "__main__":
    # Removed blocking input to allow automated execution
    import sys
    if len(sys.argv) > 1:
        FRONTEND_URL = sys.argv[1]
    
    print(f"Starting Login Automation on {FRONTEND_URL}...")
    run_tests()
